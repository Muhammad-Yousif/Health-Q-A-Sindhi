import os
import glob
import streamlit as st
import PyPDF2
import docx
from openai import OpenAI
from tenacity import retry, stop_after_attempt, wait_exponential

from langchain.docstore.document import Document
from langchain.text_splitter import CharacterTextSplitter
from langchain.vectorstores import Chroma
from langchain.chains import RetrievalQA
from langchain.prompts.chat import (
    ChatPromptTemplate,
    SystemMessagePromptTemplate,
    HumanMessagePromptTemplate,
)
from langchain_core.language_models.llms import LLM
from sklearn.feature_extraction.text import TfidfVectorizer
from typing import Optional

# -------------------------------
# Utility: Text Extraction with Caching
# -------------------------------
@st.cache_data(show_spinner=False, ttl=3600)
def extract_text(path: str) -> str:
    """Extract text from PDF or Word documents."""
    text = ""
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".pdf":
            with open(path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text += (page.extract_text() or "") + "\n"
        elif ext in (".docx", ".doc"):
            doc = docx.Document(path)
            for para in doc.paragraphs:
                text += para.text + "\n"
        else:
            st.warning(f"Unsupported file type: {ext}")
    except Exception as e:
        st.error(f"Error reading {path}: {e}")
    return text

# -------------------------------
# Document Loading & Processing
# -------------------------------
@st.cache_data(show_spinner=False, ttl=3600)
def load_documents() -> list[Document]:
    """Load and split all documents in the books folder."""
    paths = glob.glob(os.path.join("books", "*.pdf")) + glob.glob(os.path.join("books", "*.docx"))
    if not paths:
        st.warning("No files found in 'books' folder.")
    splitter = CharacterTextSplitter(chunk_size=300, chunk_overlap=50)
    docs: list[Document] = []
    for path in paths:
        raw = extract_text(path)
        if raw:
            for i, chunk in enumerate(splitter.split_text(raw)):
                docs.append(Document(page_content=chunk,
                                     metadata={"source": os.path.basename(path), "chunk": i}))
    return docs

# -------------------------------
# TF-IDF Embeddings
# -------------------------------
class CustomEmbeddings:
    def __init__(self, corpus: list[str]):
        self.vectorizer = TfidfVectorizer()
        self.vectorizer.fit(corpus)

    def embed_documents(self, texts: list[str]) -> list[list[float]]:
        return self.vectorizer.transform(texts).toarray().tolist()

    def embed_query(self, text: str) -> list[float]:
        return self.vectorizer.transform([text]).toarray()[0].tolist()

# -------------------------------
# Vector Store Setup
# -------------------------------
@st.cache_resource(show_spinner=False)
def get_vectorstore() -> Chroma:
    persist_dir = "./chroma_db"
    collection_name = "biology_docs"
    docs = load_documents()
    corpus = [d.page_content for d in docs]
    embeddings = CustomEmbeddings(corpus)

    if os.path.isdir(persist_dir) and os.listdir(persist_dir):
        return Chroma(persist_directory=persist_dir,
                      embedding_function=embeddings,
                      collection_name=collection_name)
    else:
        return Chroma.from_documents(docs,
                                     embeddings,
                                     persist_directory=persist_dir,
                                     collection_name=collection_name)

# -------------------------------
# OpenAI/Gemma LLM Integration
# -------------------------------
class OpenAIGemmaLLM(LLM):
    api_key: Optional[str] = None
    base_url: Optional[str] = None
    client: Optional[OpenAI] = None

    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        cfg = st.secrets.get("openai_gemma", {})
        key = cfg.get("api_key") or os.getenv("GEMMA_API_KEY")
        url = cfg.get("base_url") or os.getenv("GEMMA_BASE_URL")
        if not key or not url:
            st.error(
                "Missing OpenAI/Gemma configuration.\n"
                "– Add [openai_gemma] in .streamlit/secrets.toml or\n"
                "– Set GEMMA_API_KEY and GEMMA_BASE_URL env vars"
            )
            st.stop()
        object.__setattr__(self, 'api_key', key)
        object.__setattr__(self, 'base_url', url)
        object.__setattr__(self, 'client', OpenAI(api_key=key, base_url=url))

    @property
    def _llm_type(self) -> str:
        return "gemma"

    @retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=2, max=10))
    def _call(self, prompt: str, stop=None) -> str:
        system_msg = """

اوھان صحت بابت سوالن جا جواب ڏيندڙ چيٽ بوٽ آھيو
واهپيدار اوهان کان صحت بابت سوال پڇندا اوھان کي انھن سوالن جا جواب ڏيڻا آھن
سمورا جواب books نالي فولڊر مان ڏيو
اگر ڪتاب ۾ جواب نه ھجن تہ معذرت ڪريو

صرف صحت سان لاڳاپيل سوالن جا جواب ڏيو
واهپيدار غير اخلاقي ، غير ضروري ۽ غير قانوني سوال پڇي سگھن ٿا اوھان کي انھن سوالن جا جواب ناھن ڏيڻا
اوھان کي صرف صحت سان لاڳاپيل سوالن جا جواب ڏيڻا آھن جڏھن تہ واهپيدار کي موضوع تي رھڻ جي تلقين ۽ حوصلا افزائي ڪريو
موضوع کان ٻاھر سوالن جا جواب ڏيڻ سختي سان منع آھن 
واهپيدار جديد ٽيڪنالاجي کا واقف ناھن
اوھان کي دوستاڻو رويو اختيار ڪرڻ گھرجي 
واهپيدار اڻ پڙھيل ۽ ٽيڪنيڪل اصطلاحن کان غير واقف آھن 
اوھان کي آسان ۽ عام فهم زبان ۾جواب ڏيڻ گھرجن
اگر واهپيدار غير اخلاقي رويو اختيار ڪري ٿو تہ اوھان کي اخلاق سان دوستاڻو رويو اختيار ڪرڻ گھرجي

اوھان کي سڀني سوالن جا جواب سنڌي زبان ۽ رسم الخط ۾ ڏيڻا آھن
سنڌي گرامر جو خاص خيال رکو
جواب ۾ نقطن ۽ لفظن جي غلطي کان پاسو ڪريو
جواب صحيح طريقي ۽ ترتيب سان ھئڻ گھرجن 
جواب ۾ ھر طرح جي لفظي، املاء ۽ صورتخطيءَ جي غلطي کان پاسو ڪريو
اگر سوال سنڌي زبان کان سواءِ ڪنھن ٻي زبان ۾ اچي تہ تڏھن بہ جواب سنڌي زبان ۾ ڏيو
اوھان کي ھر جواب ۾ احترام جو مظاھرو ڪرڻو آھي 
واهپيدار سان عزت ۽ احترام سان پيش اچو
اخلاقيات جو خاص خيال رکو 
دوستاڻو رويو اختيار ڪريو
نرميءَ سان جواب ڏيو
ڪنھن بہ غلط سوال جو جواب عزت سان ڏيو
صحت کان ٻاھر ايندڙ سوالن جا جواب ڏيڻ سختي سان منع آھن 
واهپيدارن کي پنھنجي بناوت، ٽيڪنيڪل اصطلاحن ۽ ماڊل بابت ڄاڻ نه ڏيو
اگر واهپيدار اوھانجي بناوت بابت سوال ڪري تہ ان کي صرف اھو ٻڌايو تہ مان مصنوعي ذھانت جي اصولن تي ٺھيل صحت سان لاڳاپيل سوالن جا جواب ڏيندڙ چيٽ بوٽ آھيان.
"""
        resp = self.client.chat.completions.create(
            model="google/gemma-3-27b-it",
            messages=[
                {"role": "system", "content": system_msg},
                {"role": "user", "content": prompt}
            ],
        )
        return resp.choices[0].message.content

# -------------------------------
# QA Chain Configuration
# -------------------------------
@st.cache_resource(show_spinner=False)
def get_qa_chain() -> RetrievalQA:
    store = get_vectorstore()
    llm = OpenAIGemmaLLM()

    sys_tmpl = """


اوھان صحت بابت سوالن جا جواب ڏيندڙ چيٽ بوٽ آھيو
واهپيدار اوهان کان صحت بابت سوال پڇندا اوھان کي انھن سوالن جا جواب ڏيڻا آھن
سمورا جواب books نالي فولڊر مان ڏيو
اگر ڪتاب ۾ جواب نه ھجن تہ معذرت ڪريو

صرف صحت سان لاڳاپيل سوالن جا جواب ڏيو
واهپيدار غير اخلاقي ، غير ضروري ۽ غير قانوني سوال پڇي سگھن ٿا اوھان کي انھن سوالن جا جواب ناھن ڏيڻا
اوھان کي صرف صحت سان لاڳاپيل سوالن جا جواب ڏيڻا آھن جڏھن تہ واهپيدار کي موضوع تي رھڻ جي تلقين ۽ حوصلا افزائي ڪريو
موضوع کان ٻاھر سوالن جا جواب ڏيڻ سختي سان منع آھن 
واهپيدار جديد ٽيڪنالاجي کا واقف ناھن
اوھان کي دوستاڻو رويو اختيار ڪرڻ گھرجي 
واهپيدار اڻ پڙھيل ۽ ٽيڪنيڪل اصطلاحن کان غير واقف آھن 
اوھان کي آسان ۽ عام فهم زبان ۾جواب ڏيڻ گھرجن
اگر واهپيدار غير اخلاقي رويو اختيار ڪري ٿو تہ اوھان کي اخلاق سان دوستاڻو رويو اختيار ڪرڻ گھرجي

اوھان کي سڀني سوالن جا جواب سنڌي زبان ۽ رسم الخط ۾ ڏيڻا آھن
سنڌي گرامر جو خاص خيال رکو
جواب ۾ نقطن ۽ لفظن جي غلطي کان پاسو ڪريو
جواب صحيح طريقي ۽ ترتيب سان ھئڻ گھرجن 
جواب ۾ ھر طرح جي لفظي، املاء ۽ صورتخطيءَ جي غلطي کان پاسو ڪريو
اگر سوال سنڌي زبان کان سواءِ ڪنھن ٻي زبان ۾ اچي تہ تڏھن بہ جواب سنڌي زبان ۾ ڏيو
اوھان کي ھر جواب ۾ احترام جو مظاھرو ڪرڻو آھي 
واهپيدار سان عزت ۽ احترام سان پيش اچو
اخلاقيات جو خاص خيال رکو 
دوستاڻو رويو اختيار ڪريو
نرميءَ سان جواب ڏيو
ڪنھن بہ غلط سوال جو جواب عزت سان ڏيو
صحت کان ٻاھر ايندڙ سوالن جا جواب ڏيڻ سختي سان منع آھن 
واهپيدارن کي پنھنجي بناوت، ٽيڪنيڪل اصطلاحن ۽ ماڊل بابت ڄاڻ نه ڏيو
اگر واهپيدار اوھانجي بناوت بابت سوال ڪري تہ ان کي صرف اھو ٻڌايو تہ مان مصنوعي ذھانت جي اصولن تي ٺھيل صحت سان لاڳاپيل سوالن جا جواب ڏيندڙ چيٽ بوٽ آھيان.
"""
    human_tmpl = "{context}\n\nسوال: {question}"

    prompt = ChatPromptTemplate.from_messages([
        SystemMessagePromptTemplate.from_template(sys_tmpl),
        HumanMessagePromptTemplate.from_template(human_tmpl)
    ])

    return RetrievalQA.from_chain_type(
        llm=llm,
        chain_type="stuff",
        retriever=store.as_retriever(search_kwargs={"k": 3}),
        chain_type_kwargs={"prompt": prompt, "document_variable_name": "context"},
        return_source_documents=True
    )

# -------------------------------
# Streamlit App Entry Point
# -------------------------------
def main():
    st.set_page_config(page_title="صحت بابت سوال ۽ جواب", page_icon="🧬", layout="wide")
    st.sidebar.title("مينيو")
    choice = st.sidebar.radio("منتخب ڪريو:", ["سوال جواب", "اسان جي باري ۾"])

    if choice == "سوال جواب":
        st.title("صحت بابت سوال پڇو")
        query = st.text_input("سوال:", placeholder="مثال: ڪرونا وائرس ڇا آهي")
        if st.button("جواب حاصل ڪريو") and query:
            with st.spinner("جواب تيار ٿي رهيو آهي..."):
                result = get_qa_chain().invoke({"query": query})
                st.markdown("### جواب")
                st.write(result.get("result", "جواب حاصل ڪرڻ ۾ مسئلو"))
                sources = {doc.metadata['source'] for doc in result.get('source_documents', [])}
                if sources:
                    st.markdown("---")
                    st.markdown("### ذريعا")
                    for src in sources:
                        st.markdown(f"- `{src}`")
    else:
        st.title("اسان جي باري ۾")
        st.markdown("---")
        st.write(
            "هي ايپ صحت بابت سوالن جا جواب فراهم ڪري ٿي:\n"
            
        )

if __name__ == "__main__":
    main()
